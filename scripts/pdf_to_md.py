"""
Конвертер PDF -> Markdown (локально, БЕЗ токенов).

Берёт папку с PDF, превращает каждый в .md и (опционально) кладёт прямо в базу
знаний агента: agents/<имя>/knowledge/. Дальше агент читает их через свои
knowledge-инструменты (list/read/search_knowledge). Никакого дообучения — это RAG.

Использование:
    # просто конвертировать папку рядом (file.pdf -> file.md)
    python3 scripts/pdf_to_md.py ~/Downloads/standards

    # сразу в базу знаний агента standards_creator
    python3 scripts/pdf_to_md.py ~/Downloads/standards --agent standards_creator

    # в произвольную папку
    python3 scripts/pdf_to_md.py ~/Downloads/standards --out ~/converted

Опции:
    --agent <имя>   класть .md в agents/<имя>/knowledge/
    --out <путь>    класть .md в указанную папку
    --no-index      не создавать _index.md (оглавление базы)
"""

import sys
import re
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


def _slug(name):
    s = re.sub(r"[^\w\-. ]", "_", name).strip().replace(" ", "_")
    return re.sub(r"_+", "_", s)


def _first_heading(md, fallback):
    for line in md.splitlines():
        line = line.strip()
        if line.startswith("#"):
            return line.lstrip("# ").strip() or fallback
        if line:
            return line[:80]
    return fallback


def _pdf_to_md(path):
    import pymupdf4llm
    return pymupdf4llm.to_markdown(str(path))


def _docx_to_md(path):
    import docx
    doc = docx.Document(str(path))
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = (p.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(ch for ch in style if ch.isdigit()) or "2"
            lines.append(f"{'#' * min(int(level), 6)} {text}")
        else:
            lines.append(text)
    # таблицы — простым списком строк
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(lines)


def convert_folder(src, out_dir, make_index=True):
    src = Path(src).expanduser()
    out_dir = Path(out_dir).expanduser()
    out_dir.mkdir(parents=True, exist_ok=True)

    files = sorted([f for f in src.rglob("*")
                    if f.suffix.lower() in (".pdf", ".docx")])
    if not files:
        print(f"В {src} не найдено PDF/DOCX-файлов.")
        return []

    print(f"Найдено документов: {len(files)}. Конвертирую в {out_dir} …\n")
    index = []
    for i, doc in enumerate(files, 1):
        rel = doc.relative_to(src)
        out_name = _slug(rel.with_suffix(".md").as_posix().replace("/", "__"))
        out_path = out_dir / out_name
        try:
            md = _docx_to_md(doc) if doc.suffix.lower() == ".docx" else _pdf_to_md(doc)
            if not md.strip():
                print(f"  [{i}/{len(files)}] ПУСТО (возможно скан без текста): {doc.name}")
                continue
            out_path.write_text(md, encoding="utf-8")
            title = _first_heading(md, doc.stem)
            index.append((out_name, title, len(md)))
            print(f"  [{i}/{len(files)}] {doc.name}  ->  {out_name}  ({len(md)} симв.)")
        except Exception as e:
            print(f"  [{i}/{len(files)}] ОШИБКА {doc.name}: {e}")

    if make_index and index:
        lines = ["# База знаний — оглавление\n",
                 f"Всего документов: {len(index)}\n"]
        for fname, title, size in index:
            lines.append(f"- **{title}** — `{fname}` ({size} симв.)")
        (out_dir / "_index.md").write_text("\n".join(lines), encoding="utf-8")
        print(f"\nОглавление: {out_dir / '_index.md'}")

    print(f"\nГотово. Сконвертировано: {len(index)}/{len(files)}.")
    return index


def main():
    args = sys.argv[1:]
    if not args:
        print(__doc__)
        return

    src = args[0]
    out = None
    agent = None
    make_index = True

    i = 1
    while i < len(args):
        if args[i] == "--agent" and i + 1 < len(args):
            agent = args[i + 1]; i += 2
        elif args[i] == "--out" and i + 1 < len(args):
            out = args[i + 1]; i += 2
        elif args[i] == "--no-index":
            make_index = False; i += 1
        else:
            i += 1

    if agent:
        out = ROOT / "agents" / agent / "knowledge"
    elif not out:
        out = src  # рядом с исходниками

    convert_folder(src, out, make_index=make_index)
    if agent:
        print(f"\nФайлы в базе знаний агента «{agent}». "
              f"Он найдёт их через search_knowledge / read_knowledge.")


if __name__ == "__main__":
    main()
